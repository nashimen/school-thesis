import time, codecs, csv, math, numpy as np, random, datetime, os, gc, pandas as pd, jieba, re, sys
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.feature_extraction.text import CountVectorizer
import numpy
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from mpl_toolkits.mplot3d import Axes3D
import matplotlib.font_manager as font_manager
import docx
from docx import Document

import warnings
warnings.filterwarnings("ignore", category=Warning)

pd.set_option('display.max_columns', None)

debug = False

# 生成Top5和Topc10的检索格式
def delete_repetition(path):
    top10 = pd.read_excel(path, sheet_name="Top10")
    top5 = pd.read_excel(path, sheet_name="Top5")

    result = top10["name"].tolist()
    result = ' OR '.join(result)
    result = result.replace("  ", " ")
    print("result:", result)
    result = result.replace("and", "\"AND\"")
    print("result:", result)


# 读取Top5&Top10期刊文件
def read_top_journals(path, s_path):
    document = Document(path)
    tables = document.tables
    length = len(tables)
    result = pd.DataFrame()

    for j in range(length):
        table = tables[j]
        for i in range(len(table.rows)):
            name = table.cell(i, 1).text
            print("name = ", name)
            name = name.replace("&", "and").replace(",", "")
            result = result.append([{"name": name}])
    result.to_excel(s_path, index=False, header=None)


# 读取word文件
def read_word_fie(path, s_path):
    document = Document(path)
    tables = document.tables
    length = len(tables)
    cols = ["publisher", "field", "name", "star"]
    result = pd.DataFrame(cols)
    for j in range(length):
        # print(tables)
        print("*" * 80, "table", j, "*" * 80)
        table = tables[j]
        print("table's length:", len(table.rows))
        for i in range(len(table.rows)):
            # if i == 0:
            #     continue
            field = table.cell(i, 1).text
            name = table.cell(i, 2).text
            star = table.cell(i, 6).text
            publisher = table.cell(i, 3).text

            field = field.replace("\n", " ")
            name = name.replace("\n", " ").replace(",", "")
            star = star.replace("\n", "")

            result = result.append([{"publisher": publisher, "field": field, "name": name, "star": star}])

    result.to_excel(s_path, index=False, header=None)


if __name__ == "__main__":
    start_time = time.time()
    print("Start time : ",  time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(start_time)))

    path_top_global = "G:\期刊分级办法相关\Top5&Top10期刊.docx"
    s_path_top_global = "G:\期刊分级办法相关\ABS\\top-test.xlsx" if debug else "G:\期刊分级办法相关\ABS\\top.xlsx"
    path_global = "G:\期刊分级办法相关\ABS\AJGABS2021list.docx"
    s_path_global = "G:\期刊分级办法相关\ABS\AJGABS2021list-test.xlsx" if debug else "G:\期刊分级办法相关\ABS\AJGABS2021list.xlsx"
    # read_word_fie(path_global, s_path_global)

    # read_top_journals(path_top_global, s_path_top_global)
    delete_repetition("G:\期刊分级办法相关\ABS\AJGABS2021list.xlsx")

    end_time = time.time()
    print("End time : ",  time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(end_time)))
    total_time = end_time - start_time
    print("Consume total time:", total_time, "s")

